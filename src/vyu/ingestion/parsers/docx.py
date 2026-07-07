from __future__ import annotations

import zipfile
from dataclasses import dataclass
from io import BytesIO

from docx import Document as DocxDocument

from src.vyu.ingestion.parsers.base import (
    ParseResult,
    ParsedDocument,
    ParsedSection,
    ParsedTable,
    SourcePosition,
    extract_identifiers_from_text,
    normalize_unicode,
)

PARSER_NAME = "vyu_docx"
PARSER_VERSION = "1.0.0"
_MAX_DECOMPRESSION_RATIO = 100
_MACRO_PARTS = frozenset(
    {
        "word/vbaProject.bin",
        "word/vbaData.xml",
    }
)


@dataclass(frozen=True)
class DocxDocumentParser:
    max_decompression_ratio: int = _MAX_DECOMPRESSION_RATIO

    @property
    def name(self) -> str:
        return PARSER_NAME

    @property
    def version(self) -> str:
        return PARSER_VERSION

    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    def parse(self, data: bytes, *, filename: str, media_type: str) -> ParseResult:
        del filename, media_type
        if not data:
            return ParseResult.failed(
                parser_name=self.name,
                parser_version=self.version,
                code="malformed_document",
                message="DOCX document is empty.",
            )
        archive_error = self._validate_archive(data)
        if archive_error is not None:
            return archive_error
        try:
            document = DocxDocument(BytesIO(data))
        except (zipfile.BadZipFile, KeyError, ValueError):
            return ParseResult.failed(
                parser_name=self.name,
                parser_version=self.version,
                code="malformed_archive",
                message="DOCX archive is malformed.",
            )
        except Exception:
            return ParseResult.failed(
                parser_name=self.name,
                parser_version=self.version,
                code="malformed_document",
                message="DOCX document could not be parsed.",
            )

        sections: list[ParsedSection] = []
        tables: list[ParsedTable] = []
        warnings: list[str] = []
        current_title: str | None = None
        current_chunks: list[str] = []

        def flush_section() -> None:
            if not current_chunks and current_title is None:
                return
            text = normalize_unicode("\n".join(current_chunks).strip())
            if text or current_title:
                sections.append(
                    ParsedSection(
                        title=current_title,
                        text=text,
                        position=SourcePosition(section=current_title or "body"),
                    )
                )

        for paragraph in document.paragraphs:
            style_name = paragraph.style.name if paragraph.style is not None else ""
            text = normalize_unicode(paragraph.text.strip())
            if not text:
                continue
            if style_name.startswith("Heading"):
                flush_section()
                current_title = text
                current_chunks = []
            else:
                current_chunks.append(text)

        for table in document.tables:
            rows = tuple(
                tuple(normalize_unicode(cell.text.strip()) for cell in row.cells)
                for row in table.rows
            )
            tables.append(
                ParsedTable(
                    caption=None,
                    cells=rows,
                    position=SourcePosition(section=current_title or "table"),
                )
            )
            if not rows:
                warnings.append("unsupported_table:empty")

        flush_section()
        if not sections:
            sections.append(
                ParsedSection(
                    title=None,
                    text="",
                    position=SourcePosition(section="body"),
                )
            )

        core_props = document.core_properties
        title = normalize_unicode(core_props.title.strip()) if core_props.title else None
        if title is None and sections:
            title = sections[0].title
        authors = tuple(
            normalize_unicode(author.strip())
            for author in (core_props.author or "").split(";")
            if author.strip()
        )
        published_at = (
            core_props.modified.isoformat()
            if core_props.modified is not None
            else None
        )
        page_text = "\n\n".join(section.text for section in sections if section.text)
        identifiers = extract_identifiers_from_text(page_text)
        parsed = ParsedDocument(
            title=title,
            authors=authors,
            published_at=published_at,
            identifiers=identifiers,
            sections=tuple(sections),
            pages=(page_text,),
            tables=tuple(tables),
            figures=(),
            references=(),
            warnings=tuple(warnings),
        )
        return ParseResult.success(
            parser_name=self.name,
            parser_version=self.version,
            document=parsed,
        )

    def _validate_archive(self, data: bytes) -> ParseResult | None:
        try:
            with zipfile.ZipFile(BytesIO(data)) as archive:
                names = set(archive.namelist())
                if names & _MACRO_PARTS:
                    return ParseResult.failed(
                        parser_name=self.name,
                        parser_version=self.version,
                        code="macro_enabled_office",
                        message="Macro-enabled Office files are not supported.",
                    )
                uncompressed = sum(info.file_size for info in archive.infolist())
                compressed = max(len(data), 1)
                if uncompressed > compressed * self.max_decompression_ratio:
                    return ParseResult.failed(
                        parser_name=self.name,
                        parser_version=self.version,
                        code="excessive_compression_ratio",
                        message="DOCX archive exceeds the compression ratio limit.",
                    )
                for name in names:
                    lowered = name.lower()
                    if lowered.endswith((".exe", ".dll", ".bat", ".cmd", ".js", ".vbs")):
                        return ParseResult.failed(
                            parser_name=self.name,
                            parser_version=self.version,
                            code="embedded_executable",
                            message="DOCX archive contains embedded executable content.",
                        )
        except zipfile.BadZipFile:
            return ParseResult.failed(
                parser_name=self.name,
                parser_version=self.version,
                code="malformed_archive",
                message="DOCX archive is malformed.",
            )
        return None
