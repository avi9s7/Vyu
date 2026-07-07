from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document as DocxDocument
from pypdf import PdfWriter

from src.vyu.ingestion.malware import EICAR_TEST_SIGNATURE

PUBLIC_ARTICLE_HTML = """<!DOCTYPE html>
<html>
  <head>
    <title>Public Health Operations Review</title>
    <script>alert('xss')</script>
    <style>.hidden { display: none; }</style>
  </head>
  <body>
    <h1>Public Health Operations Review</h1>
    <p>This public health article discusses hospital operations without patient identifiers.</p>
    <table>
      <caption>Staffing Summary</caption>
      <tr><th>Unit</th><th>Headcount</th></tr>
      <tr><td>ICU</td><td>24</td></tr>
    </table>
    <figure><figcaption>Figure 1. Ward layout schematic.</figcaption></figure>
  </body>
</html>
"""

PUBLIC_ARTICLE_TEXT = """Public Health Operations Review

This public health article discusses hospital operations without patient identifiers.
DOI: 10.5555/public.article
"""

SYNTHETIC_PHI_TEXT = "Patient ID: ABC-12345. Date of birth: 01/15/1980."


def build_minimal_pdf(*, page_count: int = 1) -> bytes:
    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=612, height=792)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def build_sample_docx() -> bytes:
    document = DocxDocument()
    document.core_properties.title = "VYU Sample Report"
    document.core_properties.author = "VYU Fixtures"
    document.add_heading("VYU Sample Report", level=1)
    document.add_paragraph("Body text with DOI 10.1234/vyu.test and PMID: 12345678.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Throughput"
    table.cell(1, 1).text = "42"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_macro_docx() -> bytes:
    base = build_sample_docx()
    output = BytesIO()
    with ZipFile(BytesIO(base), "r") as source, ZipFile(output, "w", ZIP_DEFLATED) as archive:
        for name in source.namelist():
            archive.writestr(name, source.read(name))
        archive.writestr("word/vbaProject.bin", b"fake-macro")
    return output.getvalue()


@dataclass(frozen=True)
class StagingUploadFixture:
    name: str
    filename: str
    media_type: str
    body: bytes
    expect_ready: bool
    expect_code: str | None = None


CLEAN_STAGING_FIXTURES: tuple[StagingUploadFixture, ...] = (
    StagingUploadFixture(
        name="clean_txt",
        filename="report.txt",
        media_type="text/plain",
        body=PUBLIC_ARTICLE_TEXT.encode("utf-8"),
        expect_ready=True,
    ),
    StagingUploadFixture(
        name="clean_pdf",
        filename="report.pdf",
        media_type="application/pdf",
        body=build_minimal_pdf(page_count=2),
        expect_ready=True,
    ),
    StagingUploadFixture(
        name="clean_docx",
        filename="report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        body=build_sample_docx(),
        expect_ready=True,
    ),
    StagingUploadFixture(
        name="clean_html",
        filename="report.html",
        media_type="text/html",
        body=PUBLIC_ARTICLE_HTML.encode("utf-8"),
        expect_ready=True,
    ),
)

BLOCKED_STAGING_FIXTURES: tuple[StagingUploadFixture, ...] = (
    StagingUploadFixture(
        name="eicar_malware",
        filename="eicar.txt",
        media_type="text/plain",
        body=EICAR_TEST_SIGNATURE.encode("ascii"),
        expect_ready=False,
        expect_code="malware_infected",
    ),
    StagingUploadFixture(
        name="synthetic_phi",
        filename="clinical-note.txt",
        media_type="text/plain",
        body=SYNTHETIC_PHI_TEXT.encode("utf-8"),
        expect_ready=False,
        expect_code="phi_suspected_phi",
    ),
)
